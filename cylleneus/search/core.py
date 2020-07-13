import codecs
import json
import re
from datetime import datetime
from math import ceil
from pathlib import Path
from typing import Iterable

import docx
import safer

from cylleneus import __version__, settings
from cylleneus.corpus.core import Corpus, Work
from cylleneus.engine import scoring
from cylleneus.engine.highlight import (
    CylleneusBasicFragmentScorer,
    CylleneusDefaultFormatter,
    CylleneusPinpointFragmenter,
)
from cylleneus.engine.qparser.default import CylleneusQueryParser
from cylleneus.engine.searching import CylleneusSearcher, HitRef
from cylleneus.utils import Debug, slugify


class Collection:
    def __init__(self, works: Iterable[Work] = None, name: str = None):
        self._works = list(works) if works else []
        self._name = name

        if self.name:
            self.load()

    @property
    def name(self):
        return self._name

    @name.setter
    def name(self, s: str):
        self._name = s

    def add(self, work):
        if work not in self._works:
            self._works.append(work)

    def remove(self, work):
        if work in self._works:
            self._works.remove(work)

    def save(self, name: str = None):
        if not name:
            name = self.name
        else:
            self.name = name

        file = Path(".collections")

        if file.exists():
            with codecs.open(file, "r", "utf8") as fp:
                collections = json.load(fp)
        else:
            collections = {}

        collections[name] = list(self.iter_all())
        with safer.open(".collections", "w", encoding="utf8") as fp:
            json.dump(collections, fp)

    def load(self, name: str = None):
        if not name:
            name = self.name

        file = Path(".collections")
        if file.exists():
            with codecs.open(file, "r", "utf8") as fp:
                collections = json.load(fp)
        else:
            collections = {name: []}

        if name in collections:
            self.works = [
                Corpus(corpus).work_by_docix(docix)
                for corpus, docixs in collections[name]
                for docix in docixs
            ]
            self.name = name

    @property
    def count(self):
        return len(self._works)

    @property
    def works(self):
        return self._works

    @works.setter
    def works(self, w: list):
        self._works = w

    def iter_all(self):
        for work in self.works:
            yield (work.corpus.name, work.docix)

    def __iter__(self):
        yield from self.works

    def __str__(self):
        return str([str(work) for work in self.works])

    def __repr__(self):
        return f"Collection(works={self.works})"


class Searcher:
    def __init__(self, collection: Collection = None):
        self._searches = []
        self._collection = collection

    @property
    def collection(self):
        return self._collection

    @collection.setter
    def collection(self, c):
        self._collection = c

    def search(self, spec: str, minscore=None):
        """ Execute the specified search specification """

        search = Search(spec, self.collection, minscore=minscore)
        _ = search.run()
        self.searches.append(search)
        return search

    @property
    def searches(self):
        return self._searches

    @property
    def history(self):
        return self.searches


class Search:
    def __init__(
        self,
        spec: str,
        collection: Collection,
        minscore=None,
        top=1000000
    ):
        self._spec = spec
        self._collection = collection
        self._minscore = minscore
        self._top = top

        self._query = None
        self._start_dt = None
        self._end_dt = None
        self._results = None
        self._count = None
        self._highlights = None

        self._maxchars = 70  # width of one line
        self._surround = (
            70 if settings.CHARS_OF_CONTEXT < 70 else settings.CHARS_OF_CONTEXT
        )

    @classmethod
    def from_json(cls, s):
        obj = json.loads(s)
        s = Search(
            spec=obj["spec"],
            collection=Collection(
                works=[
                    Corpus(corpus).work_by_docix(docix)
                    for corpus, docix in obj["collection"]
                ]
            ),
            minscore=obj["minscore"],
            top=obj["top"],
        )
        s.start_dt = datetime.fromisoformat(obj["start_dt"])
        s.end_dt = datetime.fromisoformat(obj["end_dt"])
        s.maxchars = obj["maxchars"]
        s.surround = obj["surround"]
        s.count = obj["count"]

        hlites = []
        for r in obj["results"]:
            href = HitRef(
                r["corpus"],
                r["author"],
                r["title"],
                r["urn"],
                r["reference"],
                r["text"],
            )
            hlites.append(href)
        s.highlights = hlites
        return s

    @property
    def collection(self):
        return self._collection

    @property
    def docixs(self):
        return [
            (doc["corpus"], doc["docix"])
            for work in self.collection
            for doc in work.doc
        ]

    @property
    def docs(self):
        yield from self.collection

    @property
    def maxchars(self):
        return self._maxchars

    @maxchars.setter
    def maxchars(self, n):
        self._results = n

    @property
    def surround(self):
        return self._surround

    @surround.setter
    def surround(self, n):
        self._surround = n

    @property
    def results(self):
        if self._results is None:
            self.run()
        return self._results

    def results_to_json(self):
        if self.results:
            s = {
                "query":    self.spec,
                "start_dt": str(self.start_dt),
                "end_dt":   str(self.end_dt),
                "count":    self.count[0],
            }
            results = []
            for hit, meta, _ in self.results:
                r = {
                    "corpus": hit["corpus"],
                    "docix":  hit["docix"],
                    "author": hit["author"],
                    "title":  hit["title"],
                    "urn":    hit["urn"],
                    "meta":   {
                        "start": {
                            div: meta["start"][div]
                            for div in meta["meta"].split("-")
                        },
                        "end":   {
                            div: meta["end"][div]
                            for div in meta["meta"].split("-")
                        },
                    },
                }
                results.append(r)
            s["results"] = results
            return s

    @property
    def highlights(self):
        if not self._highlights:
            self._highlights = []
            if self.results:
                for hit, meta, fragment in self.results:
                    c = Corpus(hit["corpus"])
                    corpus, author, title, urn, reference, text = c.fetch(
                        hit, meta, fragment
                    )
                    href = HitRef(corpus, author, title, urn, reference, text)
                    self._highlights.append(href)
        yield from self._highlights

    @highlights.setter
    def highlights(self, hlites):
        self._highlights = hlites

    def to_dict(self):
        if self.results:
            s = {
                "query":      self.spec,
                "collection": [
                    {"corpus": work.corpus.name, "docix": work.docix}
                    for work in self.collection
                ],
                "minscore":   self.minscore,
                "top":        self.top,
                "start_dt":   str(self.start_dt),
                "end_dt":     str(self.end_dt),
                "maxchars":   self.maxchars,
                "surround":   self.surround,
                "count":      self.count,
            }
            results = []
            for href in self.highlights:
                r = {
                    "corpus":    href.corpus,
                    "author":    href.author,
                    "title":     href.title,
                    "urn":       href.urn,
                    "reference": href.reference,
                    "text":      href.text,
                }
                results.append(r)
            s["results"] = results
            return s

    def to_json(self):
        if self.results:
            return json.dumps(self.to_dict())

    def to_text(self):
        for href in self.highlights:
            yield href.corpus, href.author, href.title, href.urn, href.reference, href.text

    def to_docx(self, filename: str = None):
        if not filename:
            filename = slugify(self.spec, allow_unicode=False)
        else:
            if not filename.endswith(".docx"):
                filename += ".docx"

        if self.results:
            doc = docx.Document()
            doc.add_heading(self.spec, 0)
            doc.add_heading(
                f"{self.start_dt.strftime(settings.LONG_DATE_FORMAT)} (Cylleneus v{__version__})",
                2,
            )

            for corpus, author, title, urn, reference, text in self.to_text():
                h = doc.add_heading(level=1)
                h.add_run(f"{author}, ")
                r = h.add_run(f"{title} ")
                r.font.italic = True
                h.add_run(f"{reference}")

                p = doc.add_paragraph()
                for run in re.finditer(
                    r"<(\w+?)>(.*?)</\1>", text, flags=re.DOTALL
                ):
                    if run.group(1) == "match":
                        for t in run.group(2).split():
                            if re.match(r"<em>.*?</em>", t):
                                em = re.match(r"<em>(.*?)</em>", t).group(1)
                                r = p.add_run(em + " ")
                                r.font.bold = True
                            elif t.startswith("<em>"):
                                r = p.add_run(re.sub(r"<em>", "", t) + " ")
                                r.font.bold = True
                            elif t.startswith("</em>"):
                                r = p.add_run(re.sub(r"</em>", "", t) + " ")
                                r.font.bold = False
                            elif t == "\n":
                                r.add_break()
                            else:
                                r = p.add_run(t + " ")
                                r.font.bold = None
                    else:
                        p.add_run(run.group(2))
            doc.save(filename)

    @property
    def spec(self):
        return self._spec

    @results.setter
    def results(self, r):
        self._results = r

    @property
    def start_dt(self):
        return self._start_dt

    @start_dt.setter
    def start_dt(self, t):
        self._start_dt = t

    @property
    def start(self):
        return self.start_dt.strftime(settings.LONG_DATE_FORMAT)

    @property
    def end_dt(self):
        return self._end_dt

    @end_dt.setter
    def end_dt(self, t):
        self._end_dt = t

    @property
    def end(self):
        return self.end_dt.strftime(settings.LONG_DATE_FORMAT)

    @property
    def duration(self):
        if self.end_dt and self.start_dt:
            d = self.end_dt - self.start_dt
            return datetime.utcfromtimestamp(d.total_seconds()).strftime(
                settings.DURATION_FORMAT
            )

    @property
    def count(self):
        if not self._count:
            if self.results and len(self.results) > 0:
                results = [(hit, meta) for hit, meta, _ in self.results]

                corpora = len({hit["corpus"] for hit, _ in results})
                docs = len({(hit["corpus"], hit["docix"]) for hit, _ in results})
                matches = ceil(
                    sum(
                        len(
                            {tuple(hlite.values()) for hlite in meta["hlites"]}
                        )
                        for _, meta in results
                        if "hlites" in meta
                    )
                    / self.query.nterms()
                )

                self._count = matches, docs, corpora
            else:
                self._count = 0, 0, 0
        return self._count

    @count.setter
    def count(self, n):
        self._count = n  # matches, docs, corpora

    @property
    def minscore(self):
        return self._minscore

    @property
    def top(self):
        return self._top

    def run(self):
        self.start_dt = datetime.now()
        self.results = []

        for work in self.collection:
            if work.searchable:
                parser = CylleneusQueryParser("form", work.corpus.schema)
                self.query = parser.parse(self.spec)
                Debug.print(Debug.LOW, "Query: {}".format(self.query))

                for ix in work.indexes:
                    reader = ix.reader()
                    with CylleneusSearcher(
                        reader, weighting=scoring.NullWeighting
                    ) as searcher:
                        results = searcher.search(
                            self.query, terms=True, limit=None
                        )

                        if results:
                            results.fragmenter = CylleneusPinpointFragmenter(
                                autotrim=True,
                                charlimit=None,
                                maxchars=self.maxchars,
                                surround=self.surround,
                            )
                            results.scorer = CylleneusBasicFragmentScorer()
                            results.formatter = CylleneusDefaultFormatter()

                            for hit in sorted(
                                results,
                                key=lambda x: (
                                    x["corpus"],
                                    x["author"],
                                    x["title"],
                                ),
                            ):
                                self.results += hit.highlights(
                                    fieldname="content",
                                    top=self.top,
                                    minscore=self.minscore,
                                )
        self.end_dt = datetime.now()
        return self.count

    @property
    def query(self):
        return self._query

    @query.setter
    def query(self, q):
        self._query = q

    def __repr__(self):
        return f"Search(query={self.spec}, collection={self.collection})"

    def __str__(self):
        return self.spec

    def __iter__(self):
        if self.results:
            yield from self.highlights

    def __bool__(self):
        return any(self.count)
